import io


import pdfplumber
import pypdfium2 as pdfium
from docx import Document
from pdfminer.pdfdocument import PDFPasswordIncorrect

from app.core.logging import logger
from app.exceptions.custom_exceptions import (
    ParsingException,
    PasswordProtectedPDFException,
)


def extract_text_from_pdf(file_content: bytes) -> str:
    """Extracts text from a PDF file using pdfplumber with a fallback to pypdfium2.

    Args:
        file_content: The binary content of the PDF file.

    Returns:
        Extracted text as a string.

    Raises:
        PasswordProtectedPDFException: If the PDF requires a password.
        ParsingException: If text extraction fails completely or yields no text.
    """
    text_content = []

    try:
        with pdfplumber.open(io.BytesIO(file_content)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_content.append(page_text)

        extracted_text = "\n".join(text_content).strip()

        # If pdfplumber couldn't extract anything (maybe it's a weirdly formatted PDF),
        # fallback to pypdfium2
        if not extracted_text:
            logger.info("pdfplumber extracted no text. Falling back to pypdfium2.")
            pdf = pdfium.PdfDocument(file_content)
            for i in range(len(pdf)):
                page = pdf.get_page(i)
                textpage = page.get_textpage()
                text_content.append(textpage.get_text_bounded())
            extracted_text = "\n".join(text_content).strip()

        if not extracted_text:
            raise ParsingException(
                "No text could be extracted from the PDF. It might be an image-based scanned document."
            )

        return extracted_text

    except PDFPasswordIncorrect:
        raise PasswordProtectedPDFException()
    except ParsingException:
        raise
    except Exception as e:
        logger.error(f"Failed to extract text from PDF: {str(e)}")
        raise ParsingException(f"Failed to parse PDF document: {str(e)}")


def extract_text_from_docx(file_content: bytes) -> str:
    """Extracts text from a DOCX file using python-docx.

    Args:
        file_content: The binary content of the DOCX file.

    Returns:
        Extracted text as a string.

    Raises:
        ParsingException: If text extraction fails completely.
    """
    try:
        doc = Document(io.BytesIO(file_content))
        text_content = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)

        extracted_text = "\n".join(text_content).strip()

        if not extracted_text:
            raise ParsingException("No text could be extracted from the DOCX.")

        return extracted_text

    except ParsingException:
        raise
    except Exception as e:
        logger.error(f"Failed to extract text from DOCX: {str(e)}")
        raise ParsingException(f"Failed to parse DOCX document: {str(e)}")


def extract_text(file_content: bytes, extension: str) -> str:
    """Routes the file to the appropriate text extractor based on its extension.

    Args:
        file_content: The binary content of the file.
        extension: The file extension (e.g., '.pdf', '.docx').

    Returns:
        Extracted text as a string.
    """
    if extension.lower() == ".pdf":
        return extract_text_from_pdf(file_content)
    elif extension.lower() in [".docx", ".doc"]:
        return extract_text_from_docx(file_content)
    else:
        raise ParsingException(f"Unsupported extension for extraction: {extension}")
