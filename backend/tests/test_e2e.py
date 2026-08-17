import io
import os

import pytest
from app.core.config import settings
from docx import Document
from fastapi.testclient import TestClient

from app.main import app

client = TestClient(app)


@pytest.mark.skipif(not settings.GROQ_API_KEY, reason="GROQ_API_KEY is not set")
def test_e2e_analyze_flow():
    """End-to-end integration test for full ATS resume analysis API."""
    doc = Document()
    doc.add_heading("John Doe", 0)
    doc.add_paragraph("Senior Backend Engineer")
    doc.add_paragraph("Email: john.doe@example.com | Location: San Francisco, CA")
    doc.add_heading("SUMMARY", level=1)
    doc.add_paragraph(
        "Senior Software Engineer with 6+ years of experience designing high-throughput REST APIs in Python and FastAPI."
    )
    doc.add_heading("EXPERIENCE", level=1)
    doc.add_paragraph("Senior Backend Engineer at TechFlow Inc. (2021 – Present)")
    doc.add_paragraph(
        "• Architected RESTful microservices using Python, FastAPI, and Pydantic handling 15M+ daily requests."
    )
    doc.add_heading("TECHNICAL SKILLS", level=1)
    doc.add_paragraph("Python, FastAPI, SQL, Docker, Kubernetes, AWS")
    doc.add_heading("EDUCATION", level=1)
    doc.add_paragraph("Bachelor of Science in Computer Science - UC Berkeley (2018)")

    resume_stream = io.BytesIO()
    doc.save(resume_stream)
    resume_bytes = resume_stream.getvalue()

    jd_content = """Job Title: Senior Backend Engineer (Python / FastAPI)
Employment Type: Full-time

Requirements:
- 5+ years of experience in software development using Python and FastAPI.
- Deep expertise in SQL and PostgreSQL query optimization.
- Hands-on containerization experience with Docker and Kubernetes.

Responsibilities:
- Architect and deploy scalable microservices.
- Qualifications: Bachelor's degree in Computer Science."""

    temp_resume = "temp_resume.docx"
    temp_jd = "temp_jd.txt"

    with open(temp_resume, "wb") as f:
        f.write(resume_bytes)

    with open(temp_jd, "w", encoding="utf-8") as f:
        f.write(jd_content)

    try:
        with open(temp_resume, "rb") as f_resume, open(temp_jd, "rb") as f_jd:
            files = {
                "resume": (
                    "temp_resume.docx",
                    f_resume,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                ),
                "jd": ("temp_jd.txt", f_jd, "text/plain"),
            }
            response = client.post("/api/v1/analyze", files=files)

        assert response.status_code == 200
        data = response.json()
        assert "metadata" in data
        assert "match_report" in data
        assert "parsed_resume" in data
        assert "parsed_jd" in data
        assert data["match_report"]["overall_score"] > 0
    finally:
        if os.path.exists(temp_resume):
            os.remove(temp_resume)
        if os.path.exists(temp_jd):
            os.remove(temp_jd)
